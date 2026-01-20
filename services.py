"""
Servicios consolidados para el sistema de vectorización
"""
import re
import os
import asyncio
import hashlib
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
from settings import settings

# FastAPI
from fastapi import UploadFile
import aiofiles

# Procesamiento de documentos
import fitz  # PyMuPDF
import pdfplumber
from docx import Document

# Vectorización
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    PayloadSchemaType,
    Filter,
    FieldCondition,
    MatchValue
)
from qdrant_client.http.exceptions import UnexpectedResponse

from langchain.text_splitter import RecursiveCharacterTextSplitter

_vectorization_service: Optional["VectorizationService"] = None

# =============================================================================
# UTILIDADES
# =============================================================================

def generate_document_hash_from_content(text: str) -> str:
    """Genera hash SHA256 del contenido del texto"""
    normalized = re.sub(r"\s+", " ", text.strip().lower())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _sanitize_filename(filename: str) -> str:
    safe_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_."
    safe_name = "".join(c if c in safe_chars else "_" for c in filename)

    if len(safe_name) > 100:
        name_parts = safe_name.rsplit(".", 1)
        if len(name_parts) == 2:
            safe_name = name_parts[0][:90] + "." + name_parts[1]
        else:
            safe_name = safe_name[:100]

    return safe_name

# ============================================================================
# PARTE 1: CARGA DE DOCUMENTOS
# ============================================================================

async def load_and_save_document(file: UploadFile, upload_dir: Path, task_id: str) -> Path:
    """
    Guarda el archivo subido y retorna su path
    """
    # Validar extensión
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in settings.ALLOWED_EXTENSIONS:
        raise ValueError(f"Extensión {file_ext} no permitida. Use: {settings.ALLOWED_EXTENSIONS}")

    # Generar nombre único
    safe_filename = f"{task_id}_{_sanitize_filename(file.filename)}"
    file_path = upload_dir / safe_filename

    # Guardar archivo de forma asíncrona
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()

        # Validar tamaño
        if len(content) > settings.MAX_FILE_SIZE:
            raise ValueError(f"Archivo muy grande. Máximo: {settings.MAX_FILE_SIZE / 1024 / 1024:.1f} MB")

        await f.write(content)

    return file_path


# ============================================================================
# PARTE 2: VALIDACIÓN DE DOCUMENTOS
# ============================================================================

async def validate_document_structure(file_path: str) -> Dict[str, Any]:
    """
    Valida la estructura del documento y detecta elementos no textuales
    """
    path = Path(file_path)

    if path.suffix.lower() == '.pdf':
        return await _validate_pdf(file_path)
    elif path.suffix.lower() == '.docx':
        return await _validate_docx(file_path)
    else:
        raise ValueError(f"Formato no soportado: {path.suffix}")


async def _validate_pdf(file_path: str) -> Dict[str, Any]:
    """Validación específica para archivos PDF"""
    stats = {
        "total_pages": 0,
        "total_images": 0,
        "total_tables": 0,
        "pages_with_issues": [],
        "recommendation": None
    }

    try:
        with pdfplumber.open(file_path) as pdf:
            stats["total_pages"] = len(pdf.pages)

            for i, page in enumerate(pdf.pages):
                page_num = i + 1
                has_issue = False

                # Detectar tablas
                tables = page.find_tables(table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "intersection_tolerance": 5,
                    "snap_tolerance": 5
                })

                if len(tables) > 0:
                    stats["total_tables"] += len(tables)
                    has_issue = True

                # Detectar imágenes (filtrar decoraciones pequeñas)
                images = page.images
                relevant_images = [
                    img for img in images
                    if float(img.get('width', 0)) > 50 and float(img.get('height', 0)) > 50
                ]

                if len(relevant_images) > 0:
                    stats["total_images"] += len(relevant_images)
                    has_issue = True

                if has_issue:
                    stats["pages_with_issues"].append(page_num)

        # Validación adicional con PyMuPDF para imágenes
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, 1):
            img_list = page.get_images(full=True)
            real_images = [img for img in img_list if img[2] > 100 and img[3] > 100]
            if real_images and page_num not in stats["pages_with_issues"]:
                stats["total_images"] += len(real_images)
                stats["pages_with_issues"].append(page_num)
        doc.close()

    except Exception as e:
        raise Exception(f"Error al validar PDF: {str(e)}")

    # Generar recomendación
    if stats["total_images"] > 0 or stats["total_tables"] > 0:
        stats["recommendation"] = (
            f"Se detectaron {stats['total_images']} imagen(es) y {stats['total_tables']} tabla(s). "
            "Solo se vectorizará el texto del documento."
        )
    else:
        stats["recommendation"] = "El documento está limpio y listo para vectorización."

    return stats


async def _validate_docx(file_path: str) -> Dict[str, Any]:
    """Validación específica para archivos DOCX"""
    stats = {
        "total_pages": 1,
        "total_images": 0,
        "total_tables": 0,
        "pages_with_issues": [],
        "recommendation": None
    }

    try:
        doc = Document(file_path)

        # Detectar tablas
        stats["total_tables"] = len(doc.tables)

        # Detectar imágenes embebidas
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                stats["total_images"] += 1

        # Estimar páginas (aproximado: 500 palabras por página)
        total_words = sum(len(para.text.split()) for para in doc.paragraphs)
        stats["total_pages"] = max(1, total_words // 500)

    except Exception as e:
        raise Exception(f"Error al validar DOCX: {str(e)}")

    # Generar recomendación
    if stats["total_images"] > 0 or stats["total_tables"] > 0:
        stats["recommendation"] = (
            f"Se detectaron {stats['total_images']} imagen(es) y {stats['total_tables']} tabla(s). "
            "Solo se vectorizará el texto del documento."
        )
    else:
        stats["recommendation"] = "El documento está limpio y listo para vectorización."

    return stats




# ============================================================================
# PARTE 3: PROCESAMIENTO DE TEXTO
# ============================================================================

async def _extract_text(file_path: str) -> str:
    """Extrae texto de PDF o DOCX preservando estructura"""
    path = Path(file_path)

    if path.suffix.lower() == '.pdf':
        return await _extract_from_pdf(file_path)
    elif path.suffix.lower() == '.docx':
        return await _extract_from_docx(file_path)
    else:
        raise ValueError(f"Formato no soportado: {path.suffix}")


async def _extract_from_pdf(file_path: str) -> str:
    """Extracción optimizada de texto desde PDF, excluyendo tablas, encabezados y pies de página."""
    doc = fitz.open(file_path)
    full_text = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text_elements = []

                page_height = page.height
                header_bottom = page_height * 0.10
                footer_top = page_height * 0.90

                tables = page.find_tables(table_settings={
                        "vertical_strategy": "lines",
                        "horizontal_strategy": "lines",
                        "intersection_tolerance": 4,})
                table_bboxes = [t.bbox for t in tables]

                words = page.extract_words(x_tolerance=1, y_tolerance=1)
                
                for word in words:
                    word_bbox = (word['x0'], word['top'], word['x1'], word['bottom'])
                    
                    # Check if word is in header or footer region
                    is_header_footer = word['top'] < header_bottom or word['bottom'] > footer_top

                    # Check if word overlaps with any table
                    is_in_table = False
                    for t_bbox in table_bboxes:
                        # Simple overlap check: if word's bbox intersects table's bbox
                        if not (word_bbox[2] < t_bbox[0] or word_bbox[0] > t_bbox[2] or \
                                word_bbox[3] < t_bbox[1] or word_bbox[1] > t_bbox[3]):
                            is_in_table = True
                            break

                    if not is_header_footer and not is_in_table:
                        page_text_elements.append(word['text'])

                if page_text_elements:
                    #full_text.append(f"\n--- Página {page_num} ---\n")
                    full_text.append(" ".join(page_text_elements))

    except Exception as e:
        raise Exception(f"Error al extraer texto del PDF: {str(e)}")

    doc.close()
    return "\n".join(full_text)


async def _extract_from_docx(file_path: str) -> str:
    """Extracción de texto desde DOCX"""
    try:
        doc = Document(file_path)
        paragraphs = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    except Exception as e:
        raise Exception(f"Error al extraer texto del DOCX: {str(e)}")


async def _normalize_text(text: str) -> str:
    """Normalización avanzada del texto"""
    text = re.sub(r'[^\w\s.,;:()/\-%¿?!áéíóúñÁÉÍÓÚÑ]', '', text)
    text = re.sub(r'\x00', '', text)
    text = re.sub(r'[\x01-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)

    return text.strip()


async def extract_normalize_and_hash(file_path: str) -> Dict[str, Any]:
    """
    Extrae texto, normaliza y genera hash del documento
    Función pública para usar en main.py
    """
    # ✅ CORRECCIÓN: Usar funciones con guion bajo
    raw_text = await _extract_text(file_path)
    normalized_text = await _normalize_text(raw_text)

    if len(normalized_text) < 100:
        raise ValueError("Documento sin contenido suficiente")

    return {
        "normalized_text": normalized_text,
        "document_hash": generate_document_hash_from_content(normalized_text)
    }


async def _chunk_text_semantic(text: str) -> List[Dict[str, Any]]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=[
            "\n\n",
            "\n",
            ". ",
            ", ",
            " ",
            ""
        ]
    )

    results = []
    chunk_id = 0

    page_chunks = splitter.split_text(text)

    for chunk in page_chunks:
        urls = re.findall(r'https?://\S+', chunk)

        results.append({
            "chunk": chunk_id,
            "text": chunk.strip(),
            "urls": urls if urls else None
        })
        chunk_id += 1

    return results

async def chunk_text_semantic(text: str) -> List[Dict[str, Any]]:
    """
    Chunking semántico - Función pública para exportar
    """
    return await _chunk_text_semantic(text)


# ============================================================================
# PARTE 4: VECTORIZACIÓN Y ALMACENAMIENTO
# ============================================================================

class VectorizationService:
    """Servicio singleton para vectorización"""
    _instance = None
    _model = None
    _client = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self):
        if self._model is None:
            print(f"🔄 Cargando modelo: {settings.EMBEDDING_MODEL}")
            self._model = SentenceTransformer(settings.EMBEDDING_MODEL)
            print("✅ Modelo cargado")

        if self._client is None:
            print(f"🔄 Conectando a Qdrant: {settings.QDRANT_HOST}:{settings.QDRANT_PORT}")
            self._client = QdrantClient(
                url=f"http://{settings.QDRANT_HOST}:{settings.QDRANT_PORT}",
                timeout=60
            )
            print("✅ Conexión exitosa")


    def collection_exists(self, collection_name: str) -> bool:
        try:
            collections_response = self._client.get_collections()
            existing_names = [c.name for c in collections_response.collections]
            
            print(f"🔍 Nombres en Qdrant: {existing_names}")
            print(f"🔎 Buscando exactamente: '{collection_name}'")
            
            return collection_name in existing_names
        except Exception as e:
            print(f"❌ Error consultando colecciones: {e}")
            return False

    def document_exists(
        self,
        document_hash: str,
        collection_name: str
    ) -> bool:
        try:
            result, _ = self._client.scroll(
                collection_name=collection_name,
                limit=1,
                scroll_filter=Filter(
                    must=[
                        FieldCondition(
                            key="document_hash",
                            match=MatchValue(value=document_hash)
                        )
                    ]
                )
            )
            return len(result) > 0
        except Exception:
            return False

    def get_documents_in_collection(self, collection_name: str) -> List[str]:
            """
            Obtiene la lista de nombres de documentos únicos en una colección

            Returns:
                Lista de nombres de archivos únicos
            """
            try:
                # Hacer scroll de todos los puntos para obtener filenames únicos
                documents = set()
                offset = None

                while True:
                    results, next_offset = self._client.scroll(
                        collection_name=collection_name,
                        limit=100,
                        offset=offset,
                        with_payload=["filename"],
                        with_vectors=False
                    )

                    for point in results:
                        filename = point.payload.get("filename")
                        if filename:
                            documents.add(filename)

                    if next_offset is None:
                        break

                    offset = next_offset

                return list(documents)

            except Exception as e:
                print(f"❌ Error obteniendo documentos: {e}")
                return []

    def document_exists_by_filename(
                self,
                filename: str,
                collection_name: str
        ) -> Dict[str, Any]:
            """
            Verifica si un documento existe por nombre de archivo

            Returns:
                {
                    "exists": bool,
                    "total_chunks": int,
                    "document_hash": str (si existe)
                }
            """
            try:
                # Buscar puntos con ese filename
                results, _ = self._client.scroll(
                    collection_name=collection_name,
                    limit=1,
                    scroll_filter=Filter(
                        must=[
                            FieldCondition(
                                key="filename",
                                match=MatchValue(value=filename)
                            )
                        ]
                    ),
                    with_payload=True
                )

                if not results:
                    return {
                        "exists": False,
                        "total_chunks": 0,
                        "document_hash": None
                    }

                # Contar cuántos chunks tiene este documento
                document_hash = results[0].payload.get("document_hash")

                # Contar todos los chunks con este hash
                all_results, _ = self._client.scroll(
                    collection_name=collection_name,
                    limit=10000,  # Límite alto para contar todos
                    scroll_filter=Filter(
                        must=[
                            FieldCondition(
                                key="document_hash",
                                match=MatchValue(value=document_hash)
                            )
                        ]
                    ),
                    with_payload=False,
                    with_vectors=False
                )

                return {
                    "exists": True,
                    "total_chunks": len(all_results),
                    "document_hash": document_hash
                }

            except Exception as e:
                print(f"❌ Error verificando documento: {e}")
                return {
                    "exists": False,
                    "total_chunks": 0,
                    "document_hash": None
                }

    def delete_document_by_filename(
                self,
                filename: str,
                collection_name: str
        ) -> Dict[str, Any]:
            """
            Elimina TODOS los puntos asociados a un documento específico

            Args:
                filename: Nombre del archivo a eliminar
                collection_name: Colección donde está el documento

            Returns:
                Diccionario con resultado de la operación
            """
            try:
                # Primero obtener el document_hash
                check = self.document_exists_by_filename(filename, collection_name)

                if not check["exists"]:
                    return {
                        "success": False,
                        "message": f"Documento '{filename}' no existe en la colección"
                    }

                document_hash = check["document_hash"]

                print(f"🗑️ Eliminando {check['total_chunks']} chunks del documento: {filename}")

                self._client.delete(
                    collection_name=collection_name,
                    points_selector=Filter(
                        must=[
                            FieldCondition(
                                key="document_hash",
                                match=MatchValue(value=document_hash)
                            )
                        ]
                    )
                )

                print(f"✅ Documento '{filename}' eliminado exitosamente")

                return {
                    "success": True,
                    "deleted_chunks": check["total_chunks"],
                    "filename": filename,
                    "message": f"✅ {check['total_chunks']} chunks eliminados"
                }

            except Exception as e:
                print(f"❌ Error eliminando documento: {e}")
                return {
                    "success": False,
                    "message": f"Error: {str(e)}"
                }


    async def embed(
            self,
            texts: List[str],
            *,
            is_query: bool = False
        ) -> List[List[float]]:
            """
            Genera embeddings usando BGE (query / passage)
            """
            prefix = "query: " if is_query else "passage: "
            inputs = [prefix + t for t in texts]

            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(
                None,
                lambda: self._model.encode(
                    inputs,
                    normalize_embeddings=True
                ).tolist()
            )

    async def store_vectors(
            self,
            *,
            collection_name: str,
            chunks: List[Dict[str, Any]],
            embeddings: List[List[float]],
            metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Almacena vectores en Qdrant con verificación de duplicados
        """

        if not self.collection_exists(collection_name):
            raise ValueError(f"La colección '{collection_name}' no existe en Qdrant. Operación abortada.")

        document_hash = metadata["document_hash"]
        now = datetime.utcnow().isoformat()
        points = []

        for chunk, vector in zip(chunks, embeddings):
            pid = hashlib.md5(
                f"{document_hash}_{chunk['chunk']}".encode()
            ).hexdigest()

            points.append(
                PointStruct(
                    id=pid,
                    vector=vector,
                    payload={
                        "document_hash": document_hash,
                        "filename": metadata["filename"],
                        "format": metadata["format"],
                        "total_pages": metadata["total_pages"],
                        "total_chunks": metadata["total_chunks"],
                        "date": now,
                        "chunk": chunk["chunk"],
                        "urls": chunk["urls"],
                        "text": chunk["text"]
                    }
                )
            )

        print(f"💾 Insertando {len(points)} vectores en {collection_name}...")
        batch_size = 100
        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            self._client.upsert(
                collection_name=collection_name,
                points=batch,
                wait=True
            )

        return {
            "success": True,
            "document_hash": document_hash,
            "vectors_stored": len(points),
            "collection": collection_name,
            "message": f"✅ {len(points)} vectores almacenados exitosamente en {collection_name}"
        }

# =============================================================================
# FACTORY
# =============================================================================

def check_duplicate_by_filename(
    filename: str,
    collection_name: str
) -> Dict[str, Any]:
    """
    Wrapper para verificar duplicados desde main.py
    """
    service = init_vectorization_service()
    if not service.collection_exists(collection_name):
        return {"exists": False, "total_chunks": 0}
    return service.document_exists_by_filename(filename, collection_name)


def delete_document_from_collection(
    filename: str,
    collection_name: str
) -> Dict[str, Any]:
    """
    Wrapper para eliminar documento desde main.py
    """
    service = init_vectorization_service()
    return service.delete_document_by_filename(filename, collection_name)

def init_vectorization_service():
    """
    Inicializa el servicio de vectorización
    """
    global _vectorization_service
    if _vectorization_service is None:
        _vectorization_service = VectorizationService()
    return _vectorization_service